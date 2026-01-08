"""File handlers for different input formats (TXT, DOCX, MD)."""

from pathlib import Path
from typing import Union
import logging

logger = logging.getLogger(__name__)


class FileHandler:
    """Base file handler interface."""

    @staticmethod
    def read(file_path: Union[str, Path]) -> str:
        """
        Read file content.

        Args:
            file_path: Path to file

        Returns:
            File content as string

        Raises:
            NotImplementedError: Must be implemented by subclasses
        """
        raise NotImplementedError


class TxtHandler(FileHandler):
    """Handler for TXT files."""

    @staticmethod
    def read(file_path: Union[str, Path]) -> str:
        """
        Read TXT file.

        Args:
            file_path: Path to TXT file

        Returns:
            Text content
        """
        path = Path(file_path)
        logger.info(f"Reading TXT file: {path}")

        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'cp1251', 'latin-1']

        for encoding in encodings:
            try:
                content = path.read_text(encoding=encoding)
                logger.info(f"Successfully read with {encoding} encoding")
                return content.strip()
            except UnicodeDecodeError:
                continue

        # Fallback: read as bytes and decode with errors='replace'
        logger.warning("Using fallback encoding with errors='replace'")
        content = path.read_bytes().decode('utf-8', errors='replace')
        return content.strip()


class DocxHandler(FileHandler):
    """Handler for DOCX files."""

    @staticmethod
    def read(file_path: Union[str, Path]) -> str:
        """
        Read DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Text content

        Raises:
            ImportError: If python-docx not installed
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required to read DOCX files. Install with: pip install python-docx")

        path = Path(file_path)
        logger.info(f"Reading DOCX file: {path}")

        doc = Document(path)

        # Extract all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)

        content = "\n\n".join(paragraphs)
        logger.info(f"Extracted {len(paragraphs)} paragraphs from DOCX")
        return content.strip()


class MarkdownHandler(FileHandler):
    """Handler for Markdown files."""

    @staticmethod
    def read(file_path: Union[str, Path]) -> str:
        """
        Read Markdown file.

        Args:
            file_path: Path to MD file

        Returns:
            Markdown content
        """
        path = Path(file_path)
        logger.info(f"Reading Markdown file: {path}")

        content = path.read_text(encoding='utf-8')
        return content.strip()


class FileHandlerFactory:
    """Factory for getting appropriate file handler."""

    HANDLERS = {
        '.txt': TxtHandler,
        '.docx': DocxHandler,
        '.md': MarkdownHandler,
        '.markdown': MarkdownHandler,
    }

    @classmethod
    def get_handler(cls, file_path: Union[str, Path]) -> FileHandler:
        """
        Get appropriate handler for file.

        Args:
            file_path: Path to file

        Returns:
            FileHandler instance

        Raises:
            ValueError: If file format not supported
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in cls.HANDLERS:
            raise ValueError(f"Unsupported file format: {ext}. Supported: {', '.join(cls.HANDLERS.keys())}")

        return cls.HANDLERS[ext]

    @classmethod
    def read_file(cls, file_path: Union[str, Path]) -> str:
        """
        Read file using appropriate handler.

        Args:
            file_path: Path to file

        Returns:
            File content

        Raises:
            ValueError: If file format not supported
        """
        handler = cls.get_handler(file_path)
        return handler.read(file_path)


def normalize_text(text: str) -> str:
    """
    Normalize text (line breaks, spaces, etc.).

    Args:
        text: Input text

    Returns:
        Normalized text
    """
    # Normalize line breaks
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Remove excessive blank lines (more than 2 consecutive)
    import re
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Strip leading/trailing whitespace
    text = text.strip()

    return text


def save_markdown(content: str, file_path: Union[str, Path]):
    """
    Save content as Markdown file.

    Args:
        content: Markdown content
        file_path: Path to save file
    """
    path = Path(file_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    # Normalize before saving
    content = normalize_text(content)

    path.write_text(content, encoding='utf-8')
    logger.info(f"Saved Markdown file: {path}")
